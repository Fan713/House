#! /usr/bin/env python
#coding=gbk
from docx import Document
from docx.shared import Inches
import re,os
# import SendKeys
#模拟键盘Ctrl+A
# import win32api
# import win32con
# def full_cho():
#     win32api.keybd_event(17, 0, 0, 0)
#     win32api.keybd_event(65, 0, 0, 0)
#     win32api.keybd_event(17, 0, win32con.KEYEVENTF_KEYUP, 0)
#     win32api.keybd_event(65, 0, win32con.KEYEVENTF_KEYUP, 0)
#     
t1=r'C:\Users\KLJS154\Desktop\testwe.docx'
document = Document(t1)
for paragraph in document.paragraphs:
#     print(paragraph.text)
#     SendKeys.Sendkeys('A')
#     del(paragraph)
#     full_cho()
    paragraph.clear()
#插入不同等级的标题
document.add_heading(u'MS WORD写入测试',0)
document.add_heading(u'环境配置23',1)
document.add_heading(u'添加总线23',2)
#添加文本
document.add_paragraph('测试管理')
document.save(r'C:\Users\KLJS154\Desktop\testwe.docx')




